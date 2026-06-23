"""Extract plain text from resume PDF and DOCX files."""

import io
import logging
from pathlib import Path

import pdfplumber

log = logging.getLogger(__name__)

MAX_PAGES = 5
MIN_TEXT_LENGTH = 80


def _normalize_text(text: str) -> str:
    full_text = text.strip()
    if len(full_text) < MIN_TEXT_LENGTH:
        raise ValueError(
            "Could not extract text from the resume. Please upload a text-based PDF or DOCX file."
        )
    return full_text


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = pdf.pages[:MAX_PAGES]
            chunks = []
            for page in pages:
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append(text.strip())
    except Exception as error:
        raise ValueError(f"Could not read PDF: {error}") from error

    full_text = _normalize_text("\n\n".join(chunks))
    log.info("Extracted %s characters from resume PDF", len(full_text))
    return full_text


def _extract_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as error:
        raise ValueError("DOCX support is not installed on the server.") from error

    try:
        document = Document(io.BytesIO(file_bytes))
        chunks = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))
    except Exception as error:
        raise ValueError(f"Could not read DOCX: {error}") from error

    full_text = _normalize_text("\n".join(chunks))
    log.info("Extracted %s characters from resume DOCX", len(full_text))
    return full_text


def file_extension(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def extract_text(file_bytes: bytes, filename: str = "resume.pdf") -> str:
    """Return plain text from a PDF or DOCX resume."""
    if not file_bytes:
        raise ValueError("Empty file uploaded.")

    ext = file_extension(filename)
    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    if ext == ".docx":
        return _extract_docx(file_bytes)

    raise ValueError("Please upload a PDF or DOCX file only.")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Backward-compatible helper for PDF-only callers."""
    return _extract_pdf(file_bytes)
