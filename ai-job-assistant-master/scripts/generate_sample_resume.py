#!/usr/bin/env python3
"""Generate sample resume files (PDF + DOCX) for demos and tests."""

from pathlib import Path

from docx import Document
from fpdf import FPDF

DATA_DIR = Path(__file__).resolve().parents[1] / "data" / "sample_resumes"
PDF_OUTPUT = DATA_DIR / "sample_btech_resume.pdf"
DOCX_OUTPUT = DATA_DIR / "sample_btech_resume.docx"

RESUME_LINES = [
    ("heading", "Ravi Kumar"),
    ("line", "Email: ravi.kumar.demo@example.com | Phone: +91-9876543210"),
    ("line", "Hyderabad, Telangana"),
    ("section", "Objective"),
    (
        "body",
        "Motivated B.Tech graduate seeking a Python Developer role. Strong in backend "
        "development, SQL, and building academic projects with Flask.",
    ),
    ("section", "Education"),
    ("body", "B.Tech Computer Science - JNTU Hyderabad (2022-2026)\nCGPA: 8.1/10"),
    ("section", "Skills"),
    (
        "body",
        "Python, SQL, Flask, REST APIs, Pandas, NumPy, Git, HTML, CSS, JavaScript, "
        "MySQL, Basic Machine Learning, Data Structures",
    ),
    ("section", "Projects"),
    (
        "body",
        "1. Student Attendance Tracker - Flask + SQLite web app for college attendance.\n"
        "2. Expense Splitter Bot - Python script to split hostel expenses among roommates.\n"
        "3. Weather Dashboard - Fetches Open-Meteo API and shows charts with Chart.js.",
    ),
    ("section", "Internship"),
    (
        "body",
        "Python Intern - Local startup (Summer 2025)\n"
        "Built internal data export scripts and fixed bugs in Flask admin panel.",
    ),
]


def build_pdf():
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    for kind, text in RESUME_LINES:
        if kind == "heading":
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=11)
        elif kind == "line":
            pdf.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
        elif kind == "section":
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=11)
        else:
            pdf.multi_cell(0, 6, text)
            pdf.ln(1)

    # ensure output directory exists before writing PDF
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    pdf.output(str(PDF_OUTPUT))
    print(f"Wrote {PDF_OUTPUT}")


def build_docx():
    document = Document()
    for kind, text in RESUME_LINES:
        if kind == "heading":
            document.add_heading(text, level=0)
        elif kind == "section":
            document.add_heading(text, level=2)
        else:
            document.add_paragraph(text)

    DATA_DIR.mkdir(parents=True, exist_ok=True)
    document.save(str(DOCX_OUTPUT))
    print(f"Wrote {DOCX_OUTPUT}")


if __name__ == "__main__":
    build_pdf()
    build_docx()
